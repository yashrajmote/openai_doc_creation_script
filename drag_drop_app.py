#!/usr/bin/env python3
"""
Drag and Drop Desktop Application for Excel to DOCX Generator
A simple drag-and-drop interface using tkinter with file drop functionality.
"""

import tkinter as tk
from tkinter import ttk, messagebox, scrolledtext
import tkinterdnd2 as tkdnd
import os
import threading
import openai
import pandas as pd
from dotenv import load_dotenv
from excel_to_docx_generator import ExcelToDocxGenerator

# Load environment variables
load_dotenv()


class DragDropApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Excel to DOCX Generator - Drag & Drop")
        self.root.geometry("700x600")
        self.root.resizable(True, True)
        
        # Variables
        self.output_directory = tk.StringVar()
        self.resume_output_directory = tk.StringVar()
        self.processing = False
        
        # Hardcoded template paths - Update these to your actual template file paths
        self.java_template_path = "/Users/yash/Desktop/openai_doc_creation_script/templates/java_resume_template.docx"
        self.csharp_template_path = "/Users/yash/Desktop/openai_doc_creation_script/templates/csharp_resume_template.docx"
        
        # OpenAI API key from environment variable
        self.openai_api_key = os.getenv('OPENAI_API_KEY')
        if not self.openai_api_key:
            messagebox.showerror("Error", "OpenAI API key not found. Please set OPENAI_API_KEY environment variable or create a .env file.")
            self.root.quit()
            return
        
        # Resume templates
        self.java_resume = """SKILLS -	Languages: Java, C#, Python, C++, JavaScript, TypeScript -	Frameworks and Libraries: Spring Boot, Angular, React, Next.js, Tailwind CSS, Material-UI -	Databases & Tools: MySQL, PostgreSQL, MongoDB, Firebase -	DevOps and Cloud: Docker, GitHub Actions, AWS (S3, Lambda), Git, CI/CD -	Other Tools: Kafka, Jenkins, Jira, Sanity CMS, Jasmine (Unit Testing) WORK EXPERIENCE Abhitech Energycon Limited, Toledo, OH							  	 May 2024 – Dec 2024 Full Stack Software Engineer Intern •	Developed a Gain/Loss Dashboard for coal power plants using Angular and TailwindCSS, helping users identify and act on data patterns contributing to monthly operational losses. •	Built an ETL pipeline with Apache Kafka to extract data from SAP into MySQL, improved query speed by 50% with indexing. •	Implemented Docker for containerization and integrated with GitHub Actions for CI/CD pipeline. Leveraged AWS S3 to store large volumes of SAP data and AWS Lambda to automate data processing, reducing manual intervention. Crown Equipment, New Bremen, OH	           August 2023 – Dec 2023 Full Stack Software Engineer intern •	Developed an internal invoicing application using Angular, Typescript and designed RESTful APIs in Spring Boot, Java to reduce invoice processing time by 90%. Built a reusable Radio Button List component to improve form input UX across the invoice app. •	Utilized SQL queries and stored procedures to process large datasets within the invoicing platform, resulting in 50% faster financial reporting. •	Implemented Unit Tests using Jasmine to validate Angular components, identified critical edge-cases and performed debugging, achieving a 95% test coverage. •	Collaborated with cross-functional teams in the Software Development Life Cycle in an Agile/Scrum environment, performing Code Review and QA through Jira-based workflows, and maintained technical documentation. Technochrafts, remote	  January 2023 – July 2023 Software Engineer intern •	Developed a secure cross-platform registration and login system using Java and Spring Boot to create RESTful APIs. Integrated OAuth 2.0 JWT authentication using Spring Security to ensure session management. •	Designed and implemented a responsive, user-friendly webpage using HTML, CSS, and JavaScript, ensuring cross-browser compatibility. Configured and managed the web server with NGINX to enhance performance, load balancing, and server-side caching. •	Implemented a CI/CD pipeline using GitLab CI for automated testing, integration, and deployment. Deployed containerized applications to AWS via AWS Elastic Beanstalk for simplified deployment and auto scaling of app. PROJECTS ML Intern – Anonymous Insurance Company •	Built an automated MLOps using Docker and MLFlow to retrain and evaluate XGBoost models on insurance datasets. The pipeline optimized model performance, reduced tuning time by 25%, and seamlessly handled updates for datasets exceeding 10 million rows. Headstarter Fellowship – Pantry Tracker App | Next.js, React, Firebase •	Developed a web-based inventory system using Next.js and Firebase, implemented real-time updates and item categorization using Firestore listeners. Headstarter Fellowship – AI Customer Support •	Built a real-time AI-powered chat assistant using OpenAI API and Next.js, backed by AWS Lambda and WebSockets to handle 10K+ concurrent requests with <200ms latency. EDUCATION University of Toledo, Toledo, OH Bachelor of Science Degree Recipient | GPA – 3.3 | Major: Computer Science Engineering Honors and Awards: Dean's List (2020 – 2022), UToledo Rockets Scholarship, Engineering Scholarship"""

        self.csharp_resume = """YASHRAJ MOTE LinkedIn | ymote@rockets.utoledo.edu |  GitHub SKILLS -	Languages: C#, Java, Python, C++, JavaScript, TypeScript -	Frameworks and Libraries: ASP.NET Core, Entity Framework Core, Spring Boot -	Frontend: Blazor, Angular, React, Next.js, Tailwind CSS, HTML5, CSS3 -	Databases & Tools: Microsoft SQL Server, MySQL, PostgreSQL, MongoDB, Firebase -	DevOps and Cloud: Docker, Kubernetes, GitHub Actions, GitLab CI, AWS (S3, Lambda), Azure, Git, CI/CD -	Other Tools: Kafka, Jenkins, Jira, Sanity CMS, Jasmine (Unit Testing) WORK EXPERIENCE Abhitech Energycon Limited, Toledo, OH							  	 May 2024 – Dec 2024 Full Stack Software Engineer Intern •	Developed a Gain/Loss Dashboard for coal power plants using Blazor, ASP.NET Core Web API and Entity Framework Core, helping users identify and act on data patterns contributing to monthly operational losses. •	Built an ETL pipeline with Apache Kafka to extract data from SAP into Microsoft SQL Server, and improved query speed by 50% with indexing. •	Implemented Docker for containerization and integrated with GitHub Actions for CI/CD pipeline. Leveraged AWS S3 to store large volumes of SAP data and AWS Lambda to automate data processing, reducing manual intervention. Crown Equipment, New Bremen, OH	           August 2023 – Dec 2023 Full Stack Software Engineer intern •	Developed an internal invoicing application using Angular, Typescript and designed RESTful APIs in Spring Boot, Java to reduce invoice processing time by 90%. Built a reusable Radio Button List component to improve form input UX. •	Utilized SQL queries and stored procedures to process large datasets within the invoicing platform, resulting in 50% faster financial reporting. •	Implemented Unit Tests using Jasmine to validate Angular components, identified critical edge-cases and performed debugging, achieving a 95% test coverage. •	Collaborated with cross-functional teams in the Software Development Life Cycle in an Agile/Scrum environment, performing Code Review and QA through Jira-based workflows, and maintained technical documentation. Technochrafts, remote	  January 2023 – July 2023 Software Engineer intern •	Developed a secure login system using ASP.NET Core and Entity Framework Core, implementing OAuth 2.0 and JWT for authentication and session management. •	Designed and implemented RESTful APIs in the backend service. Deployed and configured NGINX on Azure App Service to enable load balancing, implement server-side caching and optimize performance by 45%. •	Designed a CI/CD pipeline with GitLab CI and Docker, deploying to Azure Kubernetes Service to automate testing and cut deployment time by 40%. PROJECTS ML Intern – Anonymous Insurance Company •	Designed an MLOps pipeline with Docker and MLflow to automate model retraining and evaluation for XGBoost on insurance datasets. The pipeline optimized model performance, reduced tuning time by 25%, and seamlessly handled updates for datasets exceeding 10 million rows. Headstarter Fellowship – Pantry Tracker App | Next.js, React, Firebase •	Developed a web-based inventory system using Next.js and Firebase, implemented real-time updates and item categorization using Firestore listeners. Headstarter Fellowship – AI Customer Support •	Built a real-time AI-powered chat assistant using OpenAI API and Next.js, backed by AWS Lambda and WebSockets to handle 10K+ concurrent requests with <200ms latency. EDUCATION University of Toledo, Toledo, OH Bachelor of Science Degree Recipient | GPA – 3.3 | Major: Computer Science Engineering Honors and Awards: Dean's List (2020 – 2022), UToledo Rockets Scholarship, Engineering Scholarship"""
        
        
        self.setup_ui()
        
    def setup_ui(self):
        """Set up the user interface."""
        # Main frame
        main_frame = ttk.Frame(self.root, padding="20")
        main_frame.pack(fill=tk.BOTH, expand=True)
        
        # Title
        title_label = ttk.Label(main_frame, text="📄 Excel to DOCX Generator", 
                               font=("Arial", 18, "bold"))
        title_label.pack(pady=(0, 20))
        
        # Template Status Section
        template_frame = ttk.LabelFrame(main_frame, text="Resume Templates", padding="10")
        template_frame.pack(fill=tk.X, pady=10)
        
        # Template status labels
        java_status = "✅ Found" if os.path.exists(self.java_template_path) else "❌ Not Found"
        csharp_status = "✅ Found" if os.path.exists(self.csharp_template_path) else "❌ Not Found"
        
        ttk.Label(template_frame, text=f"Java Template: {java_status}").pack(anchor=tk.W)
        ttk.Label(template_frame, text=f"C# Template: {csharp_status}").pack(anchor=tk.W)
        
        if not os.path.exists(self.java_template_path) or not os.path.exists(self.csharp_template_path):
            ttk.Label(template_frame, text="⚠️ Please ensure template files exist at the specified paths", 
                     foreground="red").pack(anchor=tk.W, pady=5)
        
        # Instructions
        instructions = ttk.Label(main_frame, 
                               text="Drag and drop your Excel file here\nor click to browse",
                               font=("Arial", 12),
                               foreground="gray")
        instructions.pack(pady=10)
        
        # Process button - MOVED ABOVE DROP ZONE
        self.process_button = ttk.Button(main_frame, text="🚀 Generate DOCX Files", 
                                       command=self.process_files, 
                                       state="disabled",
                                       style="Accent.TButton")
        self.process_button.pack(pady=15)
        
        # Output directory
        dir_frame = ttk.Frame(main_frame)
        dir_frame.pack(fill=tk.X, pady=5)
        
        ttk.Label(dir_frame, text="Resume Files Directory:").pack(side=tk.LEFT)
        ttk.Entry(dir_frame, textvariable=self.resume_output_directory, width=30).pack(side=tk.LEFT, padx=(5, 5))
        ttk.Button(dir_frame, text="Browse", command=self.browse_resume_output_directory).pack(side=tk.LEFT)
        
        # AI Generated Document directory
        ai_dir_frame = ttk.Frame(main_frame)
        ai_dir_frame.pack(fill=tk.X, pady=5)
        
        ttk.Label(ai_dir_frame, text="AI Generated Document Directory:").pack(side=tk.LEFT)
        ttk.Entry(ai_dir_frame, textvariable=self.output_directory, width=30).pack(side=tk.LEFT, padx=(5, 5))
        ttk.Button(ai_dir_frame, text="Browse", command=self.browse_output_directory).pack(side=tk.LEFT)
        
        # Drop zone
        self.drop_zone = tk.Frame(main_frame, 
                                 bg="lightgray", 
                                 relief="ridge", 
                                 bd=2,
                                 height=120)
        self.drop_zone.pack(fill=tk.BOTH, expand=True, pady=10)
        self.drop_zone.pack_propagate(False)
        
        # Drop zone label
        self.drop_label = ttk.Label(self.drop_zone, 
                                   text="📁 Drop Excel file here",
                                   font=("Arial", 14),
                                   background="lightgray")
        self.drop_label.pack(expand=True)
        
        # File browser button inside drop zone
        browse_button = ttk.Button(self.drop_zone, 
                                  text="📂 Or click to browse",
                                  command=self.browse_excel_file,
                                  style="Accent.TButton")
        browse_button.pack(pady=10)
        
        # Configure drag and drop
        self.drop_zone.drop_target_register(tkdnd.DND_FILES)
        self.drop_zone.dnd_bind('<<Drop>>', self.on_drop)
        
        # Status
        self.status_label = ttk.Label(main_frame, text="Ready - Drop an Excel file to begin")
        self.status_label.pack()
        
        # Results area
        self.results_text = scrolledtext.ScrolledText(main_frame, height=8, width=60)
        self.results_text.pack(fill=tk.BOTH, expand=True, pady=10)
        
        # Set default output directories
        self.output_directory.set("/Users/yash/Desktop/Desktop - YASH's MacBook Air/APPLICATIONS")
        self.resume_output_directory.set("/Users/yash/Desktop/Desktop - YASH's MacBook Air/APPLICATIONS")
        
        # Store current file
        self.current_file = None
        
    def on_drop(self, event):
        """Handle file drop event."""
        files = self.root.tk.splitlist(event.data)
        if files:
            file_path = files[0]
            if file_path.lower().endswith(('.xlsx', '.xls')):
                self.current_file = file_path
                self.drop_label.config(text=f"✅ {os.path.basename(file_path)}")
                self.process_button.config(state="normal")
                self.status_label.config(text="File ready - Click 'Generate DOCX Files' to process")
            else:
                messagebox.showerror("Error", "Please drop an Excel file (.xlsx or .xls)")
                
    def browse_excel_file(self):
        """Browse for Excel file."""
        from tkinter import filedialog
        filename = filedialog.askopenfilename(
            title="Select Excel File",
            filetypes=[("Excel files", "*.xlsx *.xls"), ("All files", "*.*")]
        )
        if filename:
            self.current_file = filename
            self.drop_label.config(text=f"✅ {os.path.basename(filename)}")
            self.process_button.config(state="normal")
            self.status_label.config(text="File ready - Click 'Generate DOCX Files' to process")
            
    def browse_output_directory(self):
        """Browse for output directory."""
        from tkinter import filedialog
        directory = filedialog.askdirectory(title="Select DOCX Output Directory")
        if directory:
            self.output_directory.set(directory)
            
    def browse_resume_output_directory(self):
        """Browse for resume output directory."""
        from tkinter import filedialog
        directory = filedialog.askdirectory(title="Select Resume Output Directory")
        if directory:
            self.resume_output_directory.set(directory)
            
    
    def select_template(self, job_description, position):
        """Select template based on simple keyword detection."""
        text = f"{job_description} {position}".lower()
        
        # Simple keyword detection - look for C# specific terms first
        csharp_indicators = ['c#', 'csharp', '.net', 'asp.net', 'blazor', 'entity framework', 'azure']
        if any(indicator in text for indicator in csharp_indicators):
            return 'csharp'
        
        # Default to Java for everything else
        return 'java'
    
    def abbreviate_job_title(self, job_title):
        """
        Abbreviate job title according to specified rules.
        
        Args:
            job_title (str): Original job title
            
        Returns:
            str: Abbreviated job title
        """
        if not isinstance(job_title, str):
            return job_title
        
        title_lower = job_title.lower().strip()
        
        # Check for engineer -> SWE
        if 'engineer' in title_lower:
            if 'associate' in title_lower:
                return 'AssSWE'
            else:
                return 'SWE'
        
        # Check for developer -> SWD
        elif 'developer' in title_lower:
            if 'associate' in title_lower:
                return 'AssSWD'
            else:
                return 'SWD'
        
        # For other titles, return as-is (cleaned)
        import re
        invalid_chars = '<>:"/\\|?*'
        for char in invalid_chars:
            job_title = job_title.replace(char, '_')
        job_title = re.sub(r'\s+', '_', job_title.strip())
        return job_title[:50]
            
    def process_files(self):
        """Process the Excel file and generate DOCX files."""
        if not self.current_file:
            messagebox.showerror("Error", "Please drop an Excel file first.")
            return
            
        if not self.output_directory.get():
            messagebox.showerror("Error", "Please select an output directory.")
            return
            
        if not os.path.exists(self.current_file):
            messagebox.showerror("Error", "Excel file does not exist.")
            return
            
        # Start processing in a separate thread
        self.processing = True
        self.process_button.config(state="disabled")
        self.status_label.config(text="Processing...")
        self.results_text.delete(1.0, tk.END)
        
        # Run processing in separate thread
        thread = threading.Thread(target=self.run_processing)
        thread.daemon = True
        thread.start()
        
    def run_processing(self):
        """Run the actual processing in a separate thread."""
        try:
            # Check if templates exist
            if not os.path.exists(self.java_template_path) or not os.path.exists(self.csharp_template_path):
                self.root.after(0, self.show_error, "Template files not found. Please ensure both Java and C# templates exist.")
                return
            
            # Process with templates
            self.root.after(0, self.process_with_templates_only)
                
        except Exception as e:
            self.root.after(0, self.show_error, f"Error: {str(e)}")
        finally:
            self.root.after(0, self.processing_complete)
    
    def process_with_templates_only(self):
        """Process Excel file and create template-based files only."""
        try:
            df = pd.read_excel(self.current_file, header=12)
            df.columns = df.columns.str.strip()
            
            # Find the job description column
            job_desc_col = None
            for col in df.columns:
                col_lower = col.lower().strip()
                if 'description' in col_lower or 'job description' in col_lower or 'job_desc' in col_lower:
                    job_desc_col = col
                    break
            
            if not job_desc_col:
                self.results_text.insert(tk.END, "\n⚠️ No job description column found. Skipping resume generation.\n")
                return
            
            # Generate individual resume files using templates
            self.generate_individual_resumes_with_templates(None, df, job_desc_col)
            
            # Generate master document with AI responses
            self.results_text.insert(tk.END, "\n📄 Generating AI master document...\n")
            self.generate_ai_master_document(df, job_desc_col)
            
            # Update results display
            valid_entries = 0
            for index, row in df.iterrows():
                company = row.get('Company', '')
                position = row.get('Position', '')
                job_description = row.get(job_desc_col, '')
                
                if not pd.isna(company) and not pd.isna(position) and not pd.isna(job_description) and str(job_description).strip():
                    valid_entries += 1
            
            self.results_text.insert(tk.END, f"\n🎉 PROCESSING COMPLETE!\n")
            self.results_text.insert(tk.END, "=" * 50 + "\n")
            self.results_text.insert(tk.END, f"📊 Total rows processed: {len(df)}\n")
            self.results_text.insert(tk.END, f"✅ Valid entries: {valid_entries}\n")
            self.results_text.insert(tk.END, f"📄 Template files created: {valid_entries}\n")
            self.results_text.insert(tk.END, f"📁 Output directory: {self.resume_output_directory.get()}\n")
            
        except Exception as e:
            self.results_text.insert(tk.END, f"\n❌ Error in template processing: {str(e)}\n")
    
    def generate_ai_master_document(self, df, job_desc_col):
        """Generate AI master document with actual OpenAI API calls."""
        try:
            # Set up OpenAI client
            client = openai.OpenAI(api_key=self.openai_api_key)
            
            from docx import Document
            from docx.shared import Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            doc = Document()
            doc.add_heading('AI-Generated Tailored Resumes', 0)
            
            # Process each valid entry
            valid_entries = 0
            total_entries = len(df)
            
            self.results_text.insert(tk.END, f"\n🔍 Found {total_entries} total rows. Processing job descriptions...\n")
            self.results_text.see(tk.END)
            self.root.update()
            
            for index, row in df.iterrows():
                company = row.get('Company', '')
                position = row.get('Position', '')
                job_description = row.get(job_desc_col, '')
                
                # Skip if any required field is empty
                if pd.isna(company) or pd.isna(position) or pd.isna(job_description) or not str(job_description).strip():
                    continue
                
                valid_entries += 1
                self.results_text.insert(tk.END, f"\n🤖 [{valid_entries}] Generating AI resume for {company}...\n")
                self.results_text.see(tk.END)
                self.root.update()
                
                try:
                    # Generate resume using OpenAI (AI will choose template)
                    resume_content = self.generate_single_resume(client, company, position, str(job_description))
                    
                    # Add to document
                    doc.add_heading(f'{company} - {position}', level=1)
                    doc.add_paragraph(resume_content)
                    doc.add_page_break()
                    
                    self.results_text.insert(tk.END, f"✅ Completed {company}\n")
                    self.results_text.see(tk.END)
                    self.root.update()
                    
                except Exception as e:
                    self.results_text.insert(tk.END, f"❌ Error generating AI resume for {company}: {str(e)}\n")
                    continue
            
            # Save the resume document
            resume_file_path = os.path.join(self.output_directory.get(), "AI_Generated_Resumes.docx")
            doc.save(resume_file_path)
            
            self.results_text.insert(tk.END, f"\n✅ Generated {valid_entries} AI tailored resumes!\n")
            self.results_text.insert(tk.END, f"📄 AI document saved: {resume_file_path}\n")
            
        except Exception as e:
            self.results_text.insert(tk.END, f"\n❌ Error in AI master document generation: {str(e)}\n")
            
    
    def generate_individual_resumes_with_templates(self, client, df, job_desc_col):
        """Generate individual resume files by copying templates."""
        valid_entries = 0
        total_entries = len(df)
        
        self.results_text.insert(tk.END, f"\n🔍 Found {total_entries} total rows. Creating individual resume files...\n")
        self.results_text.see(tk.END)
        self.root.update()
        
        for index, row in df.iterrows():
            company = row.get('Company', '')
            position = row.get('Position', '')
            job_description = row.get(job_desc_col, '')
            
            # Skip if any required field is empty
            if pd.isna(company) or pd.isna(position) or pd.isna(job_description) or not str(job_description).strip():
                continue
            
            valid_entries += 1
            self.results_text.insert(tk.END, f"\n📄 [{valid_entries}] Creating resume file for {company}...\n")
            self.results_text.see(tk.END)
            self.root.update()
            
            try:
                # Select which template to use
                template_type = self.select_template(str(job_description), str(position))
                template_path = self.java_template_path if template_type == 'java' else self.csharp_template_path
                
                self.results_text.insert(tk.END, f"📋 Using {template_type.upper()} template\n")
                self.results_text.see(tk.END)
                self.root.update()
                
                # Load template and save as new document (proper .docx handling)
                import re
                from docx import Document
                
                safe_company = re.sub(r'[<>:"/\\|?*]', '_', str(company))[:30]
                abbreviated_position = self.abbreviate_job_title(str(position))
                resume_filename = f"{safe_company}_{abbreviated_position}.docx"
                resume_filepath = os.path.join(self.resume_output_directory.get(), resume_filename)
                
                # Load the .docx template and save as a new document
                template_doc = Document(template_path)
                template_doc.save(resume_filepath)
                
                self.results_text.insert(tk.END, f"✅ Created: {resume_filename}\n")
                self.results_text.see(tk.END)
                self.root.update()
                
            except Exception as e:
                self.results_text.insert(tk.END, f"❌ Error creating resume for {company}: {str(e)}\n")
                continue
        
        self.results_text.insert(tk.END, f"\n✅ Created {valid_entries} individual resume files!\n")
    
            
    def generate_single_resume(self, client, company, position, job_description):
        """Generate a single tailored resume using OpenAI for the master document."""
        
        prompt = f"""You are the top resume writer in the world. Your job is to take the job description I provide and tailor my resume so that it is ATS-optimized, keyword-rich, and highly compelling. Follow these steps carefully:

0. Resume Selection
   * I will provide you with two base resumes: one focused on Java roles and one focused on C# roles.
   * First, analyze the job description and decide which base resume (Java or C#) is most appropriate for this role.
   * Clearly state which base resume you selected and why in one sentence.
   * Use only the selected resume for tailoring in the following steps.

1. Keyword Extraction
   * Identify and list the most important hard skills, technical tools, industry terms, and role-specific keywords from the job description.
   * Clearly highlight which ones are must-have ATS keywords that I absolutely need in my resume.

2. Resume Tailoring
   * Rewrite my past work experience into 3–5 bullet points per role.
   * Use the XYZ method (Accomplished [X] as measured by [Y], by doing [Z]).
   * Incorporate the identified keywords naturally into the bullet points, not just in the skills section.
   * Ensure every bullet emphasizes impact, metrics, and outcomes (not just duties).
   * Update job titles if needed to better align with industry-standard titles and the target job.
   * Do not include an objective statement.

3. Keyword Integration Check
   * After writing the resume, show me exactly which keywords from the job description you integrated and where they appear (skills section, each work experience, etc.).
   * If there are important keywords you could not include, explain why.
   
Note: 
ATS Alignment
   * Ensure formatting and phrasing are ATS-friendly.
   * Avoid personal pronouns, vague buzzwords, or filler text.
   * Prioritize strong action verbs and quantified results.
---

Here are the two base resumes:
[Java Resume] - 
{self.java_resume}

[C# Resume] - 
{self.csharp_resume}

Here is the target job description:
Company: {company}
Position: {position}
Description: {job_description}"""

        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "user", "content": prompt}
            ],
            max_tokens=4000,
            temperature=0.7
        )
        
        return response.choices[0].message.content
            
        
    def show_error(self, message):
        """Show error message."""
        self.results_text.delete(1.0, tk.END)
        self.results_text.insert(tk.END, f"❌ ERROR: {message}")
        self.status_label.config(text="❌ Processing failed")
        
    def processing_complete(self):
        """Called when processing is complete."""
        self.processing = False
        self.process_button.config(state="normal")


def main():
    """Main function to run the drag and drop application."""
    root = tkdnd.Tk()
    
    # Configure style
    style = ttk.Style()
    style.theme_use('clam')
    
    # Create and run the application
    app = DragDropApp(root)
    
    # Center the window
    root.update_idletasks()
    x = (root.winfo_screenwidth() // 2) - (root.winfo_width() // 2)
    y = (root.winfo_screenheight() // 2) - (root.winfo_height() // 2)
    root.geometry(f"+{x}+{y}")
    
    root.mainloop()


if __name__ == "__main__":
    main()
