#!/usr/bin/env python3
"""
Excel to DOCX Generator Script

This script processes an Excel file with Company and Position columns,
filters out invalid entries (dates, job boards, empty cells),
and generates individual DOCX files for each valid company.
"""

import pandas as pd
import re
import os
from datetime import datetime
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import argparse
import sys


class ExcelToDocxGenerator:
    def __init__(self, excel_file_path, output_directory):
        """
        Initialize the generator with Excel file path and output directory.
        
        Args:
            excel_file_path (str): Path to the Excel file
            output_directory (str): Directory where DOCX files will be saved
        """
        self.excel_file_path = excel_file_path
        self.output_directory = output_directory
        self.job_boards = {'linkedin', 'indeed', 'handshake', 'glassdoor', 'jobright', 'github'}
        
        # Create output directory if it doesn't exist
        os.makedirs(output_directory, exist_ok=True)
    
    def is_date(self, text):
        """
        Check if the text contains a date pattern.
        
        Args:
            text (str): Text to check
            
        Returns:
            bool: True if text appears to be a date
        """
        if not isinstance(text, str):
            return False
        
        # Common date patterns
        date_patterns = [
            r'\d{1,2}-[A-Za-z]{3}',  # 10-Aug, 11-Aug
            r'\d{1,2}/\d{1,2}/\d{2,4}',  # 10/08/2024
            r'\d{1,2}-\d{1,2}-\d{2,4}',  # 10-08-2024
            r'[A-Za-z]{3}\s+\d{1,2}',  # Aug 10
            r'\d{1,2}\s+[A-Za-z]{3}',  # 10 Aug
        ]
        
        for pattern in date_patterns:
            if re.search(pattern, text, re.IGNORECASE):
                return True
        return False
    
    def is_job_board(self, text):
        """
        Check if the text is a job board name.
        
        Args:
            text (str): Text to check
            
        Returns:
            bool: True if text is a job board
        """
        if not isinstance(text, str):
            return False
        
        return text.lower().strip() in self.job_boards
    
    def is_valid_entry(self, company, position):
        """
        Check if the entry is valid for processing.
        
        Args:
            company (str): Company name
            position (str): Job position
            
        Returns:
            bool: True if entry should be processed
        """
        # Check for empty or None values
        if pd.isna(company) or pd.isna(position) or not str(company).strip() or not str(position).strip():
            return False
        
        company_str = str(company).strip()
        position_str = str(position).strip()
        
        # Check if company is a date
        if self.is_date(company_str):
            return False
        
        # Check if company is a job board
        if self.is_job_board(company_str):
            return False
        
        return True
    
    def abbreviate_job_title(self, job_title):
        """
        Abbreviate job title according to specified rules.
        
        Args:
            job_title (str): Original job title
            
        Returns:
            str: Abbreviated job title
        """
        if not isinstance(job_title, str):
            return job_title
        
        title_lower = job_title.lower().strip()
        
        # Check for engineer -> SWE
        if 'engineer' in title_lower:
            if 'associate' in title_lower:
                return 'AssSWE'
            else:
                return 'SWE'
        
        # Check for developer -> SWD
        elif 'developer' in title_lower:
            if 'associate' in title_lower:
                return 'AssSWD'
            else:
                return 'SWD'
        
        # For other titles, return as-is (cleaned)
        return self.clean_filename(job_title)
    
    def clean_filename(self, filename):
        """
        Clean filename to be safe for filesystem.
        
        Args:
            filename (str): Original filename
            
        Returns:
            str: Cleaned filename
        """
        if not isinstance(filename, str):
            return str(filename)
        
        # Remove or replace invalid characters
        invalid_chars = '<>:"/\\|?*'
        for char in invalid_chars:
            filename = filename.replace(char, '_')
        
        # Remove extra spaces and limit length
        filename = re.sub(r'\s+', '_', filename.strip())
        return filename[:50]  # Limit filename length for better readability
    
    def create_docx_file(self, company, position, index):
        """
        Create a DOCX file for a company and position.
        
        Args:
            company (str): Company name
            position (str): Job position
            index (int): Index for unique filename
            
        Returns:
            str: Path to created file
        """
        # Create new document
        doc = Document()
        
        # Add title
        title = doc.add_heading(f'Job Application: {company}', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add company information
        doc.add_heading('Company Information', level=1)
        company_para = doc.add_paragraph()
        company_para.add_run(f'Company: ').bold = True
        company_para.add_run(company)
        
        # Add position information
        doc.add_heading('Position Details', level=1)
        position_para = doc.add_paragraph()
        position_para.add_run(f'Position: ').bold = True
        position_para.add_run(position)
        
        # Add application date
        doc.add_heading('Application Details', level=1)
        date_para = doc.add_paragraph()
        date_para.add_run(f'Application Date: ').bold = True
        date_para.add_run(datetime.now().strftime('%Y-%m-%d'))
        
        # Add notes section
        doc.add_heading('Notes', level=1)
        doc.add_paragraph('Add your notes and follow-up information here...')
        
        # Generate filename using new format: CompanyName_JobTitle.docx
        clean_company = self.clean_filename(company)
        abbreviated_position = self.abbreviate_job_title(position)
        filename = f"{clean_company}_{abbreviated_position}.docx"
        filepath = os.path.join(self.output_directory, filename)
        
        # Save document
        doc.save(filepath)
        return filepath
    
    def process_excel_file(self):
        """
        Process the Excel file and generate DOCX files.
        
        Returns:
            dict: Summary of processing results
        """
        try:
            # Read Excel file starting from row 13 (0-indexed, so row 12)
            print(f"Reading Excel file: {self.excel_file_path}")
            df = pd.read_excel(self.excel_file_path, header=12)  # Row 13 is index 12
            
            # Check if required columns exist (case-insensitive)
            df.columns = df.columns.str.strip()  # Remove any whitespace
            company_col = None
            position_col = None
            
            for col in df.columns:
                if col.lower() == 'company':
                    company_col = col
                elif col.lower() == 'position':
                    position_col = col
            
            if not company_col or not position_col:
                raise ValueError(f"Excel file must contain 'Company' and 'Position' columns. Found columns: {df.columns.tolist()}")
            
            # Rename columns to standard names for easier processing
            df = df.rename(columns={company_col: 'Company', position_col: 'Position'})
            
            print(f"Found {len(df)} rows in Excel file")
            
            # Process each row
            valid_entries = 0
            skipped_entries = 0
            created_files = []
            
            for index, row in df.iterrows():
                company = row['Company']
                position = row['Position']
                
                if self.is_valid_entry(company, position):
                    try:
                        filepath = self.create_docx_file(company, position, valid_entries + 1)
                        created_files.append(filepath)
                        valid_entries += 1
                        print(f"✓ Created: {os.path.basename(filepath)}")
                    except Exception as e:
                        print(f"✗ Error creating file for {company}: {str(e)}")
                        skipped_entries += 1
                else:
                    skipped_entries += 1
                    reason = "empty cells" if pd.isna(company) or pd.isna(position) else \
                            "date in company field" if self.is_date(str(company)) else \
                            "job board name" if self.is_job_board(str(company)) else "invalid entry"
                    print(f"✗ Skipped row {index + 1}: {reason}")
            
            return {
                'total_rows': len(df),
                'valid_entries': valid_entries,
                'skipped_entries': skipped_entries,
                'created_files': created_files
            }
            
        except Exception as e:
            print(f"Error processing Excel file: {str(e)}")
            return None


def main():
    """Main function to run the script."""
    parser = argparse.ArgumentParser(description='Generate DOCX files from Excel company/position data')
    parser.add_argument('excel_file', help='Path to the Excel file')
    parser.add_argument('-o', '--output', default='./generated_docx_files', 
                       help='Output directory for DOCX files (default: ./generated_docx_files)')
    
    args = parser.parse_args()
    
    # Check if Excel file exists
    if not os.path.exists(args.excel_file):
        print(f"Error: Excel file '{args.excel_file}' not found")
        sys.exit(1)
    
    # Create generator and process file
    generator = ExcelToDocxGenerator(args.excel_file, args.output)
    results = generator.process_excel_file()
    
    if results:
        print("\n" + "="*50)
        print("PROCESSING SUMMARY")
        print("="*50)
        print(f"Total rows processed: {results['total_rows']}")
        print(f"Valid entries: {results['valid_entries']}")
        print(f"Skipped entries: {results['skipped_entries']}")
        print(f"DOCX files created: {len(results['created_files'])}")
        print(f"Output directory: {args.output}")
        
        if results['created_files']:
            print("\nCreated files:")
            for filepath in results['created_files']:
                print(f"  - {os.path.basename(filepath)}")
    else:
        print("Processing failed. Please check the error messages above.")
        sys.exit(1)


if __name__ == "__main__":
    main()
